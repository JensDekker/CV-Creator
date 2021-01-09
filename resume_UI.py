import docx
import utils
import resume

doc = "resume_template.docx"
cv = docx.Document(doc)
layout = cv.tables
print(layout)
for table in layout:
    main_col = table.cell(1, 0)
    side_col = table.cell(1, 1)

section_loc = {"Work Experience": main_col,
               "Projects": main_col,
               "Volunteer Experience": main_col,
               "Education": main_col,
               "Mechanical Design": side_col,
               "Programming": side_col,
               "Prototyping": side_col,
               "Related Courses": side_col,
               "Skills": side_col,
               "Certificates": side_col,
               "Awards": side_col,
               "Languages": side_col,
               "Interests": side_col}

data = utils.retrieve_data("resume_data.txt")

save_name = input("What's the name of the company that you're applying to? ") + "-Jens_Dekker.docx"

main_col_secs = []
side_col_secs = []
for section in data:
    if section_loc[section] == main_col:
        main_col_secs.append(section)
    else:
        side_col_secs.append(section)

q1 = "Which sections would you like to include, and what in what order? "

user_sections = utils.option_regulator(main_col_secs, q1, single_return=False)
main_col_len = len(user_sections)
user_sections = user_sections + utils.option_regulator(side_col_secs, q1, single_return=False)
side_col_len = len(user_sections) - main_col_len

user_data = {}

main_idx = 0
side_idx = 0

for section in user_sections:
    section_loc[section].add_paragraph(text="{}".format(section), style="Heading 1")

    related_course = section == "Related Courses"
    certificate = section == "Certificates"
    award = section == "Awards"

    if section_loc[section] == main_col:
        main_idx += 1
        entries = [entry for entry in data[section] if entry != "Format"]
        entry_select = utils.option_regulator(entries, "Which entries would you like to add? ", single_return=False)
        sort_entry_select = resume.date_org(entry_select, data[section])
        user_data[section] = sort_entry_select
        entry_idx = 0
        for entry in sort_entry_select:
            entry_idx += 1
            resume.class_data[section](data[section][entry]).add(section_loc[section])
            if entry_idx != len(sort_entry_select):
                main_col.add_paragraph(text="", style="Quote")
        if main_idx != main_col_len:
            main_col.add_paragraph(text="\n", style="Quote")

    elif certificate or award:
        side_idx += 1
        entries = [entry for entry in data[section] if entry != "Format"]
        entry_select = utils.option_regulator(entries, "Which entries would you like to add? ", single_return=False)
        user_data[section] = entry_select
        entry_idx = 0
        for entry in entry_select:
            entry_idx += 1
            resume.class_data[section](data[section][entry]).add(section_loc[section])
            if entry_idx != len(entry_select):
                side_col.add_paragraph(text="", style="Quote")
        if side_idx != side_col_len:
            side_col.add_paragraph(text="\n", style="Quote")

    elif related_course:
        side_idx += 1
        user_data[section] = data[section]
        resume.class_data[section](data[section]).add(section_loc[section], spaces=True)
        if side_idx != side_col_len:
            side_col.add_paragraph(text="\n", style="Quote")

    else:
        side_idx += 1
        user_data[section] = data[section]
        resume.class_data[section](data[section]).add(section_loc[section])
        if side_idx != side_col_len:
            side_col.add_paragraph(text="\n", style="Quote")


cv.save('C:\\Users\\jenst\\Desktop\\Resumes\\S2021\\' + save_name.replace(" ", "_"))
print("\nThanks for using the program!\nBest of luck with the application!")
