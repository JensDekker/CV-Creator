import docx
from docx.shared import Inches


class MainCol:
    def __init__(self, entry_dict):
        self.descr = [entry_dict["Description 1/5"],
                      entry_dict["Description 2/5"],
                      entry_dict["Description 3/5"],
                      entry_dict["Description 4/5"],
                      entry_dict["Description 5/5"]]


class Exp(MainCol):
    def __init__(self, entry_dict):
        try:
            self.group = entry_dict["Company"]
        except KeyError:
            self.group = entry_dict["Organization"]
        self.position = entry_dict["Position"]
        self.start_date = entry_dict["Start Date"]
        self.stop_date = entry_dict["Stop Date"]
        super().__init__(entry_dict)

    def add(self, column):
        dates = "{} - {}".format(self.start_date, self.stop_date)
        column.add_paragraph(text=dates, style="Normal")
        column.add_paragraph(text="{}, {}".format(self.position, self.group), style="Heading 5")
        for descr in self.descr:
            if descr != "":
                column.add_paragraph(text=descr, style="List Paragraph")


class Project(MainCol):
    def __init__(self, entry_dict):
        self.title = entry_dict["Title"]
        self.start_date = entry_dict["Start Date"]
        self.finish_date = entry_dict["Finish Date"]
        super().__init__(entry_dict)

    def add(self, column):
        dates = "{} - {}".format(self.start_date, self.finish_date)
        column.add_paragraph(text=dates, style="Normal")
        column.add_paragraph(text=self.title, style="Heading 5")
        for descr in self.descr:
            if descr != "":
                column.add_paragraph(text=descr, style="List Paragraph")


class Edu:
    def __init__(self, entry_dict):
        self.paper = entry_dict["Degree/Diploma"]
        self.institution = entry_dict["Institution"]
        self.start_date = entry_dict["Start Date"]
        self.stop_date = entry_dict["Stop Date"]
        self.GPA = entry_dict["GPA"]

    def add(self, column):
        dates = "{} - {}".format(self.start_date, self.stop_date)
        column.add_paragraph(text=dates, style="Normal")
        column.add_paragraph(text="{}, {}, GPA: {}".format(self.paper, self.institution, self.GPA), style="Normal")


class SimpleList:
    def __init__(self, entry_list):
        self.list = entry_list

    def add(self, column, spaces=False):
        if len(self.list) != 0:
            i = 0
            for item in self.list:
                i += 1
                line = column.add_paragraph(text=item, style="Normal")
                line_format = line.paragraph_format
                line_format.left_indent = Inches(0.18)
                end_of_list = i != len(self.list)
                if spaces and end_of_list:
                    column.add_paragraph(text="", style="Quote")
        else:
            print("There's nothing in the list.")


class Certificate:
    def __init__(self, entry_dict):
        self.title = entry_dict["Title"]
        self.institution = entry_dict["Institution"]

    def add(self, column):
        title = column.add_paragraph(text=self.title, style="Normal")
        title_format = title.paragraph_format
        title_format.left_indent = Inches(0.18)
        institution = column.add_paragraph(text="— {}".format(self.institution), style="Normal")
        institution_format = institution.paragraph_format
        institution_format.left_indent = Inches(0.18)


class Award:
    def __init__(self, entry_dict):
        self.title = entry_dict["Title"]
        self.organization = entry_dict["Organization"]

    def add(self, column):
        title = column.add_paragraph(text=self.title, style="Normal")
        title_format = title.paragraph_format
        title_format.left_indent = Inches(0.18)
        organization = column.add_paragraph(text="— {}".format(self.organization), style="Normal")
        organization_format = organization.paragraph_format
        organization_format.left_indent = Inches(0.18)


month_dic = {"Jan": 1,
             "Feb": 2,
             "Mar": 3,
             "Apr": 4,
             "May": 5,
             "Jun": 6,
             "Jul": 7,
             "Aug": 8,
             "Sep": 9,
             "Oct": 10,
             "Nov": 11,
             "Dec": 12}


def date_value(date):
    split_date = date.split(" ")
    return int(split_date[1]) * 15 + month_dic[split_date[0]]


def date_org(entry_lst, data_section):
    output_dic = {}
    for entry in entry_lst:
        start_date = data_section[entry]["Start Date"]
        output_dic[start_date] = entry
    dates = list(output_dic.keys())
    dates.sort(key=date_value, reverse=True)
    return [output_dic[date] for date in dates]


class_data = {"Work Experience": Exp,
              "Projects": Project,
              "Volunteer Experience": Exp,
              "Education": Edu,
              "Mechanical Design": SimpleList,
              "Programming": SimpleList,
              "Prototyping": SimpleList,
              "Skills": SimpleList,
              "Related Courses": SimpleList,
              "Certificates": Certificate,
              "Awards": Award,
              "Languages": SimpleList,
              "Interests": SimpleList}
